import io
from pathlib import Path
from typing import List, Dict, Any
import pymupdf


def load_pdf_pages(file_path: str) -> List[Dict[str, Any]]:
    """
    Extract text page-by-page from a PDF file.
    """
    pdf_path = Path(file_path)
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF not found: {file_path}")

    pages = []
    with pymupdf.open(pdf_path) as document:
        for page_idx, page in enumerate(document, start=1):
            page_text = page.get_text()
            if page_text.strip():
                pages.append(
                    {
                        "page": page_idx,
                        "text": page_text.strip(),
                    }
                )
    return pages


def load_csv_pages(file_path: str, rows_per_page: int = 20) -> List[Dict[str, Any]]:
    """
    Extract structured text from a CSV file, grouping rows into page-like chunks.
    """
    import pandas as pd

    csv_path = Path(file_path)
    if not csv_path.exists():
        raise FileNotFoundError(f"CSV not found: {file_path}")

    try:
        df = pd.read_csv(csv_path)
    except Exception:
        df = pd.read_csv(csv_path, encoding="latin-1")

    df = df.fillna("")
    total_rows = len(df)
    if total_rows == 0:
        return [{"page": 1, "text": f"Empty CSV table: {csv_path.name}"}]

    columns = list(df.columns)
    pages = []
    page_num = 1

    for start_idx in range(0, total_rows, rows_per_page):
        batch = df.iloc[start_idx : start_idx + rows_per_page]
        rows_text = []
        rows_text.append(f"Table Data: {csv_path.name} (Rows {start_idx + 1} to {min(start_idx + rows_per_page, total_rows)})")
        rows_text.append(f"Columns: {', '.join(str(c) for c in columns)}\n")

        for r_idx, row in batch.iterrows():
            row_repr = ", ".join([f"{col}: {row[col]}" for col in columns if str(row[col]).strip()])
            rows_text.append(f"Row {r_idx + 1}: {row_repr}")

        page_content = "\n".join(rows_text)
        pages.append(
            {
                "page": page_num,
                "text": page_content.strip(),
            }
        )
        page_num += 1

    return pages


def load_excel_pages(file_path: str, rows_per_page: int = 20) -> List[Dict[str, Any]]:
    """
    Extract structured text from Excel spreadsheets (.xlsx, .xls) across all sheets.
    """
    import pandas as pd

    excel_path = Path(file_path)
    if not excel_path.exists():
        raise FileNotFoundError(f"Excel file not found: {file_path}")

    xls = pd.ExcelFile(excel_path)
    pages = []
    page_num = 1

    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name).fillna("")
        total_rows = len(df)
        if total_rows == 0:
            continue

        columns = list(df.columns)
        for start_idx in range(0, total_rows, rows_per_page):
            batch = df.iloc[start_idx : start_idx + rows_per_page]
            rows_text = []
            rows_text.append(f"Sheet: {sheet_name} | Rows {start_idx + 1} to {min(start_idx + rows_per_page, total_rows)}")
            rows_text.append(f"Columns: {', '.join(str(c) for c in columns)}\n")

            for r_idx, row in batch.iterrows():
                row_repr = ", ".join([f"{col}: {row[col]}" for col in columns if str(row[col]).strip()])
                rows_text.append(f"Row {r_idx + 1}: {row_repr}")

            page_content = "\n".join(rows_text)
            pages.append(
                {
                    "page": page_num,
                    "text": page_content.strip(),
                }
            )
            page_num += 1

    if not pages:
        return [{"page": 1, "text": f"Excel Spreadsheet: {excel_path.name} (No rows found)"}]

    return pages


def load_docx_pages(file_path: str, paragraphs_per_page: int = 10) -> List[Dict[str, Any]]:
    """
    Extract text and tables from Word documents (.docx).
    """
    import docx

    doc_path = Path(file_path)
    if not doc_path.exists():
        raise FileNotFoundError(f"Word document not found: {file_path}")

    doc = docx.Document(doc_path)
    text_blocks = []

    # Extract all paragraphs
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            text_blocks.append(txt)

    # Extract all tables
    for table_idx, table in enumerate(doc.tables, 1):
        table_rows = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            if any(row_cells):
                table_rows.append(" | ".join(row_cells))
        if table_rows:
            text_blocks.append(f"\n[Table {table_idx}]\n" + "\n".join(table_rows))

    if not text_blocks:
        return [{"page": 1, "text": f"Document: {doc_path.name} (Empty)"}]

    pages = []
    page_num = 1
    for i in range(0, len(text_blocks), paragraphs_per_page):
        chunk = "\n\n".join(text_blocks[i : i + paragraphs_per_page])
        pages.append(
            {
                "page": page_num,
                "text": chunk.strip(),
            }
        )
        page_num += 1

    return pages


def load_text_pages(file_path: str, max_chars_per_page: int = 1500) -> List[Dict[str, Any]]:
    """
    Extract text from plain text files (.txt, .md).
    """
    txt_path = Path(file_path)
    if not txt_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    try:
        content = txt_path.read_text(encoding="utf-8")
    except Exception:
        content = txt_path.read_text(encoding="latin-1")

    if not content.strip():
        return [{"page": 1, "text": f"Document: {txt_path.name} (Empty)"}]

    # Split by double newlines or character limits
    paragraphs = content.split("\n\n")
    pages = []
    current_page = []
    current_len = 0
    page_num = 1

    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        if current_len + len(p) > max_chars_per_page and current_page:
            pages.append({"page": page_num, "text": "\n\n".join(current_page)})
            page_num += 1
            current_page = [p]
            current_len = len(p)
        else:
            current_page.append(p)
            current_len += len(p)

    if current_page:
        pages.append({"page": page_num, "text": "\n\n".join(current_page)})

    return pages


def load_image_pages(file_path: str) -> List[Dict[str, Any]]:
    """
    Extract text from image files (.png, .jpg, .jpeg, .webp, .tiff, .bmp) using OCR.
    """
    img_path = Path(file_path)
    if not img_path.exists():
        raise FileNotFoundError(f"Image not found: {file_path}")

    extracted_lines = []
    # 1. Try macOS native Apple Vision OCR
    try:
        from ocrmac import ocrmac
        annotations = ocrmac.OCR(str(img_path)).recognize()
        extracted_lines = [text.strip() for text, conf, bbox in annotations if text.strip()]
    except Exception:
        # 2. Fallback to pytesseract or PIL if available
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(img_path)
            raw_text = pytesseract.image_to_string(img)
            extracted_lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        except Exception:
            extracted_lines = [f"Image file: {img_path.name} (OCR could not extract text)"]

    if not extracted_lines:
        return [{"page": 1, "text": f"Image file: {img_path.name} (No readable text found in image)"}]

    full_text = "\n".join(extracted_lines)
    return [
        {
            "page": 1,
            "text": f"Scanned Document / Policy Infographic: {img_path.name}\n\n{full_text}",
        }
    ]


def load_document_pages(file_path: str) -> List[Dict[str, Any]]:
    """
    Universal multi-format document loader.
    Supports .pdf, .csv, .xlsx, .xls, .docx, .txt, .md, .png, .jpg, .jpeg, .webp, .tiff, .bmp.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf_pages(file_path)
    elif suffix == ".csv":
        return load_csv_pages(file_path)
    elif suffix in [".xlsx", ".xls"]:
        return load_excel_pages(file_path)
    elif suffix in [".docx", ".doc"]:
        return load_docx_pages(file_path)
    elif suffix in [".txt", ".md"]:
        return load_text_pages(file_path)
    elif suffix in [".png", ".jpg", ".jpeg", ".webp", ".tiff", ".bmp"]:
        return load_image_pages(file_path)
    else:
        raise ValueError(
            f"Unsupported file format '{suffix}'. "
            f"Supported formats: .pdf, .csv, .xlsx, .xls, .docx, .txt, .md, .png, .jpg, .jpeg, .webp, .tiff, .bmp"
        )


def load_pdf(file_path: str) -> str:
    """Extract combined text from all pages of a document."""
    pages = load_document_pages(file_path)
    return "\n\n".join(p["text"] for p in pages)